import re
import os
from typing import Dict, Any
import PyPDF2
import docx


def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text


def extract_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    patterns = [
        r"\+?[\d\s\-\(\)]{10,15}",
        r"\b\d{10}\b",
        r"\+\d{1,3}[\s\-]?\d{3,5}[\s\-]?\d{3,5}[\s\-]?\d{3,5}"
    ]
    for p in patterns:
        match = re.search(p, text)
        if match:
            phone = re.sub(r"[^\d\+]", "", match.group(0))
            if len(phone) >= 7:
                return match.group(0).strip()
    return ""


def extract_name_simple(text: str) -> str:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:5]:
        if 2 <= len(line.split()) <= 5 and not any(c in line for c in ["@", ":", "/"]):
            if not any(kw in line.lower() for kw in ["resume", "cv", "curriculum", "objective", "summary", "profile"]):
                return line
    return lines[0] if lines else "Unknown"


SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "react", "angular", "vue",
    "node.js", "nodejs", "express", "django", "flask", "fastapi", "spring",
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "docker", "kubernetes", "aws", "azure", "gcp", "ci/cd", "git", "github",
    "machine learning", "deep learning", "tensorflow", "pytorch", "scikit-learn",
    "data science", "pandas", "numpy", "nlp", "computer vision", "llm",
    "html", "css", "rest", "graphql", "microservices", "devops", "linux",
    "c++", "c#", ".net", "php", "ruby", "go", "rust", "swift", "kotlin",
    "agile", "scrum", "jira", "selenium", "jenkins", "terraform", "ansible",
    "excel", "tableau", "power bi", "spark", "hadoop", "kafka",
    "figma", "ui/ux", "photoshop", "illustrator",
]


def extract_skills_basic(text: str) -> list:
    text_lower = text.lower()
    found = []
    for skill in SKILL_KEYWORDS:
        if skill.lower() in text_lower:
            found.append(skill)
    return list(set(found))


def extract_experience_section(text: str) -> list:
    exp = []
    patterns = [
        r"(?i)(experience|work history|employment)(.*?)(?=education|skills|projects|certifications|$)",
    ]
    for p in patterns:
        match = re.search(p, text, re.DOTALL)
        if match:
            section = match.group(2).strip()
            lines = [l.strip() for l in section.split("\n") if l.strip()]
            if lines:
                exp.append({"raw": "\n".join(lines[:20])})
            break
    if not exp:
        exp.append({"raw": "See resume for details"})
    return exp


def extract_education_section(text: str) -> list:
    edu = []
    patterns = [
        r"(?i)(education|academic)(.*?)(?=experience|skills|projects|certifications|$)",
    ]
    for p in patterns:
        match = re.search(p, text, re.DOTALL)
        if match:
            section = match.group(2).strip()
            lines = [l.strip() for l in section.split("\n") if l.strip()]
            if lines:
                edu.append({"raw": "\n".join(lines[:10])})
            break
    if not edu:
        edu.append({"raw": "See resume for details"})
    return edu


def parse_resume(file_path: str, file_type: str) -> Dict[str, Any]:
    if file_type == "pdf":
        text = extract_text_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        text = extract_text_from_docx(file_path)
    else:
        text = ""

    return {
        "raw_text": text,
        "name": extract_name_simple(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills_basic(text),
        "work_experience": extract_experience_section(text),
        "education": extract_education_section(text),
    }
